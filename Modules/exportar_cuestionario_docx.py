from docx import Document
from docx.shared import Pt

def exportar_cuestionario_docx(ruta_markdown,ruta_docx):
    documento=Document()
    estilo=documento.styles["Normal"]
    estilo.font.name="Open sans"
    estilo.font.size=Pt(11)

    documento.add_heading("Cuestionario",level=1)

    with open(ruta_markdown,"r",encoding="utf-8") as f:
        contenido=f.readlines()
    for linea in contenido:
        linea=linea.rstrip()
        documento.add_paragraph(linea)
    documento.save(ruta_docx)

